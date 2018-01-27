#!/usr/bin/env python
# -*- coding: utf-8 -*-
#citirea si parsarea textului din manual(a documentelor word
#impartirea taskurilor pe threaduri
#Author:Tabusca Bogdan


import docx
import glob
import time
from multiprocessing import Queue
import threading
from threading import Thread
from .relation_detector import get_relations

results=[]
def set_manual(folder_path):
    files_to_search_in = glob.glob(folder_path+"\\*.docx")
    return files_to_search_in
#------------------------------------------------------------
#Get some text from those word files
def getText(filename):
    doc = docx.Document(filename)
    fullTexts = []
    no_paragraphs=len(doc.paragraphs)
    page=[]
    for i in range(1,no_paragraphs):
        if(i%100!=0):
            page.append(doc.paragraphs[i].text)
        else:
            fullTexts.append("\n".join(page))
            page=[]
            page.append(doc.paragraphs[i].text)
    return fullTexts
#-------------------------------------------------------------
# nu uita txt.encode('utf-8')
def get_Texts_for_Threads(def_folder_path):
    files_w_definition = set_manual(def_folder_path)
    texts = []
    for file in files_w_definition:
        file.encode('UTF-8')
        texts.append(getText(file))
    return texts
#------------------------------------------------------------
def do_something(term_1,term_2,text,indx):
    #do something here
    start_time = time.time()
    data=get_relations(term_1,term_2,text)
    #print(time.time()-start_time)
    results[indx]=data
    return data
#------------------------------------------------------------
#some_text=getText("E:\manual\\3_VASCULAR_2012_corr.docx")

def extract_relations(termeni_1,termeni_2,def_folder_path):
    texts=get_Texts_for_Threads(def_folder_path)
    index=0
    threads=[]
    for txt in texts:
        results.append("")
        st_txt="".join(txt)
        t=Thread(target=do_something,args=(termeni_1,termeni_2,st_txt.encode('UTF-8'),index))
        t.start()
        threads.append(t)
        index+=1
    for process in threads:
        process.join()
    final_results=[]
    for rez in results:
        for tup in rez:
            final_results.append(tup)
    final_results = set(final_results)
    final_results = list(final_results)
    return final_results

#print(extract_relations(["pacient","Paralizia","compresiunii","intracraniene","sanguin cerebral","craniu","cerebral","Răspunsul"],["tensiune","artere","criteriu","trunchi","vasculare","vase","trunchiul cerebral "],"E:\manual"))
#print(extract_relations(["Caine","Pisica","Mixer","Vietuitoare"],["Electric","Animal","vasculare"], "C:\\Users\\George\\Documents\\GitHub\\AI_OntologyMerge\Limbaj"))