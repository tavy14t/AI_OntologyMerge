import re
from lxml import html
from unidecode import unidecode
import requests
import json
from urllib.request import urlopen
from bs4 import BeautifulSoup
import re
import glob
import docx
import unicodedata
from googletrans import Translator
import requests
from bs4 import BeautifulSoup
from threading import Thread

results = []


def get_propozitii(text):
    try:
        text = text.decode('utf-8')
    except AttributeError:
        pass
    list_of_prop = re.split(";|\\.|\\!|\\?|\\n", text)
    return list_of_prop
# -------------------------------------------------


def get_good_prop(propozitii, termeni_1, termeni_2):
    good_propozitions = []
    for prop in propozitii:
        for term1 in termeni_1:
            for term2 in termeni_2:
                if term1 and term2 in prop:
                    good_propozitions.append(prop)
    good_propozitions = set(good_propozitions)
    return list(good_propozitions)
# --------------------------------------------------


def include_relation(term1, term2, prop):
    filter_1 = ["include", " are ", "detine ", "cuprinde ", "înglobează ",
                "definesc ", "defineste ", "includ ", "comasează ", "acoperă "]
    try:
        indx_1 = re.search(term1, prop).span()[1]
        prop = prop[indx_1:]
    except Exception:
        return []
    for fi_1 in filter_1:
        try:

            index_fi = re.search(fi_1, prop).span()[1]
            index_fi -= 1
            aux = prop[index_fi:]
            if re.search(term2, aux):
                return [(term1, "includes", term2)]
        except Exception:
            pass
    return []


def part_of_relation(term1, term2, prop):

    filter_1 = ["este ", "numește", "descrie", "sunt", "descriu", "numesc",
                "inclus", " incluse ", "aparține",
                "face parte", "aparțin", "fac parte"]
    filter_2 = ["o ", "un ", "din ", "în"]
    try:
        indx_1 = re.search(term1, prop).span()[1]
        prop = prop[indx_1:]
    except Exception:
        return []
    for fi_1 in filter_1:
        try:

            index_fi = re.search(fi_1, prop).span()[1]
            index_fi -= 1
            aux = prop[index_fi:]
            for fi_2 in filter_2:
                if re.search(fi_2 + term2, aux) or \
                        re.search(term2 + fi_2, aux):
                    return [(term1, "is_included", term2)]
        except Exception:
            pass
    return []


def get_relations(li_term_1, li_term_2, text):
    propozitii = get_propozitii(text)
    good_texts = get_good_prop(propozitii, li_term_1, li_term_2)
    relations = []
    for prop in good_texts:
        for t1 in li_term_1:
            for t2 in li_term_2:
                if t1 and t2 in prop:
                    relations += part_of_relation(t1, t2, prop)
                    relations += include_relation(t1, t2, prop)
    return relations


"""
print(get_relations(["Caine", "Pisica", "Mixer", "Vietuitoare"],
                    ["Electric", "Animal", "vasculare"],
                    "Pasarea este un Animal.O vietate Moarta este Inghetata."
                    "Regnul Animal include specii precum Pisica si Canine."
                    "Mixer este un Electric.Caine este un Animal."
                    "Se zice ca Pisica face parte din Animal."
                    "Clasa Vietuitoarelor include Animal."))
"""


def access_sinonim_site(word):
    try:
        with urlopen('https://dexonline.net/sinonime-%s' % word) as response:
            html = response.read()
            return html
    except Exception:
        return 'bad'


def extract_sinonims(html):
    soup = BeautifulSoup(html, 'html.parser')
    try:
        raw_sinonims = soup.find(
            "div", attrs={"class": "tip-definitie"}).find('p')
        sinonims = raw_sinonims.getText(separator=u'+')
        sinonims = unicodedata.normalize(
            'NFKD', sinonims).encode('ascii', 'ignore')
    except Exception:
        sinonims = ""
    return sinonims


def proccess_sinonims(sinonims):
    sinonims = sinonims.decode('UTF-8')
    sino = re.sub("[\(\[].*?[\)\]]", "", sinonims)
    sino = re.sub(" ", "", sino)
    sino = re.sub("s.", "", sino, 1)
    sino = re.sub("pl.", "", sino, 1)
    sino = re.sub("\.", "", sino)
    sino = re.sub("[;0-9=]", "", sino)
    processed_sinonims = []
    sino = sino.split("+")
    for x in sino:
        if x is not "":
            y = x.split(",")
            processed_sinonims = processed_sinonims + y
    # del cuv original
    try:
        processed_sinonims.pop(0)
    except Exception:
        pass
    # ========
    good_sinonims = []
    for x in processed_sinonims:
        if x is not "" and x is not "v":
            good_sinonims.append(x)
    return good_sinonims


def get_sinonim_list(l1, l2):
    l3 = []
    for word in l1:
        site = access_sinonim_site(word)
        if site is not "bad":
            raw_sino = extract_sinonims(site)
            sino = proccess_sinonims(raw_sino)
            for word2 in l2:
                for i in range(0, len(sino)):
                    if word2 == sino[i]:
                        l3.append((word, "is_synonymous", word2))
    return l3


def set_manual(folder_path):
    files_to_search_in = glob.glob("%s\\*.docx" % folder_path)
    return files_to_search_in
# -----------------------------------------------------------
# Get some text from those word files


def getText(filename):
    doc = docx.Document(filename)
    fullTexts = []
    no_paragraphs = len(doc.paragraphs)
    page = []
    for i in range(1, no_paragraphs):
        if(i % 100 != 0):
            page.append(doc.paragraphs[i].text)
        else:
            fullTexts.append("\n".join(page))
            page = []
            page.append(doc.paragraphs[i].text)
    return fullTexts
# ------------------------------------------------------------
# nu uita txt.encode('utf-8')


def get_Texts_for_Threads(def_folder_path):
    files_w_definition = set_manual(def_folder_path)
    texts = []
    for file in files_w_definition:
        file.encode('UTF-8')
        texts.append(getText(file))
    return texts
# -----------------------------------------------------------


def do_something(term_1, term_2, text, indx):
    # do something here
    # start_time = time.time()
    data = get_relations(term_1, term_2, text)
    # print(time.time()-start_time)
    results[indx] = data
    return data
# -----------------------------------------------------------
# some_text=getText("E:\manual\\3_VASCULAR_2012_corr.docx")


def extract_relations(termeni_1, termeni_2, def_folder_path):
    texts = get_Texts_for_Threads(def_folder_path)
    index = 0
    threads = []
    for txt in texts:
        results.append("")
        st_txt = "".join(txt)
        t = Thread(target=do_something, args=(
            termeni_1, termeni_2, st_txt.encode('UTF-8'), index))
        t.start()
        threads.append(t)
        index += 1
    for process in threads:
        process.join()
    final_results = []
    for rez in results:
        for tup in rez:
            final_results.append(tup)
    final_results = set(final_results)
    final_results = list(final_results)
    return final_results


def compute_name_score(str1, str2):
    # Names shorter than 3 letters are ignored
    if len(str1) < 3 or len(str2) < 3:
        return 0

    # Working with lowercase characters only
    str1 = str1.lower()
    str2 = str2.lower()

    # Split composite-words (if there are any) and calculate max score
    str1 = str1.replace('-', ' ')
    str2 = str2.replace('-', ' ')
    words1 = str1.split()
    words2 = str2.split()

    max_score = 0
    for word1 in words1:
        for word2 in words2:
            score = letter_match_length(word1, word2)
            if score > max_score:
                max_score = score

    equal_len_score = 0
    if len(str1) == len(str2):
        equal_len_score = 20
    elif max_score > 1:
        equal_len_score = 10
    return max_score**2 + equal_len_score


def compute_definition_score(str1, str2):
    def1 = get_definition(str1)
    def2 = get_definition(str2)
    def1 = def1.split()
    def2 = def2.split()
    def1 = remove_morph_words(def1)
    def2 = remove_morph_words(def2)

    reliability = 0
    for x in def1:
        for y in def2:
            if x == y:
                reliability = reliability + 1

    strength = reliability / min(len(def1), len(def2))
    return strength * reliability


def letter_match_length(str1, str2):
    i = len(str1) - 1
    j = len(str2) - 1
    length = 0
    while i >= 0 and j >= 0:
        if str1[i] == str2[j]:
            length = length + 1
        else:
            return length
        i = i - 1
        j = j - 1
    return length


def get_definition(term):
    translator = Translator()
    translation = translator.translate(term, dest='en')
    translation = remove_morph_words(translation.text.split())

    response = requests.get(
        'https://www.vocabulary.com/dictionary/' + translation[0])
    soup = BeautifulSoup(response.content, 'html.parser')

    try:
        definition = soup.find('h3', attrs={'class': 'definition'})
        definition = definition.text.strip()
        definition = ' '.join(definition[1:].split())
    except Exception:
        return '[NONE]'

    output = ''
    for c in definition:
        if c.isalpha() or c.isspace():
            output = output + c
    return output


def remove_morph_words(wordlist):
    for i in range(0, len(wordlist)):
        if wordlist[i].endswith('es'):
            wordlist[i] = wordlist[i][:-2]
        elif wordlist[i].endswith('ds') or wordlist[i].endswith('ts'):
            wordlist[i] = wordlist[i][:-1]

    output = []
    for word in wordlist:
        if word != 'a' and word != 'by' and \
            word != 'the' and word != 'as' and word != 'an'\
                and word != 'is' and word != 'of' and \
                word != 'in' and word != 'and':
            output.append(word)
    return output


def get_match_score(str1, str2):
    namescore = compute_name_score(str1, str2)
    defscore = compute_definition_score(str1, str2)
    if namescore == 0:
        namescore = 1
    if defscore == 0:
        defscore = 1
    return namescore**(1 / 2) * defscore * 2


def get_synonymus(l1, l2):
    result = []
    for i in l1:
        for j in l2:
            if (get_match_score(i, j)) >= 12:
                result.append((i, "is_synonymous", j))
    return result


def relations(l1, l2):
    s = extract_relations(
        l1, l2,
        "C:\\Users\\George\\Documents\\GitHub\\AI_OntologyMerge\\Limbaj")
    s = s + get_sinonim_list(l1, l2)
    s = s + get_synonymus(l1, l2)

    return (list(set(s)))


def get_definitions(word):
    DEX_BASE_URL = 'https://dexonline.ro'
    DEX_API_URL_FORMAT = '{}/{}'.format(DEX_BASE_URL, 'definitie/{}/json')

    dex_api_url = DEX_API_URL_FORMAT.format(word)
    dex_api_request = requests.get(dex_api_url)

    dex_raw_response = dex_api_request.json()
    dex_raw_definitions = dex_raw_response['definitions']
    dex_definitions = []

    for dex_raw_definition in dex_raw_definitions:
        dex_definition_html_rep = dex_raw_definition['htmlRep']
        html_fragments = html.fragments_fromstring(dex_definition_html_rep)
        root = html.Element('root')
        for html_fragment in html_fragments:
            root.append(html_fragment)

        dex_definition_text = root.text_content()
        dex_definition_text = unidecode(dex_definition_text)
        dex_definitions.append(dex_definition_text)
    return(dex_definitions[0].split(';')[0])


"""
l1 = ["Caine", "prieten", "Pisica", "briceag",
      "Mixer", "Vietuitoare", "fericit"]
l2 = ["Electric", "norocit", "Animal", "amic", "vasculare", "brisca"]

print(relations(l1, l2))
"""
